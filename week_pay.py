import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm
import PySimpleGUI as sg

DRIVERS = [
    "Ivan Hernandez",
    "Frederick Pena",
    "Junior A Pena",
    "Hector Vasquez",
    "Victorino Marte Perez",
    "Eligio Antonio Caba-Caba",
    "Juan Luis Acosta Mercado",
    "Joel Walter Siutto",
    "Miguel A Asencio",
    "George Berroa"
]

MAX_TRIPS = 15  # Maximum number of trips

def calculate_pay(data):
    trips = []
    for trip in data['trips']:
        if trip['trip']:
            try:
                trip_value = float(trip['trip'])
                trips.append(trip_value)
            except ValueError:
                sg.popup("Invalid input: Only numeric values can be inputted.")
                return None
        else:
            trips.append(0.0)

    try:
        diesel = float(data['diesel']) if data['diesel'] else 0.0
        tolls = float(data['tolls']) if data['tolls'] else 0.0
        detention_input = data.get('detention')
        if detention_input.strip():
            detention = float(detention_input)
        else:
            detention = 0.0
    except ValueError:
        sg.popup("Invalid input: Only numeric values can be inputted for detention time.")
        return None

    total_trips = sum(trips)
    total_expenses = diesel + tolls
    total_pay = (total_trips - total_expenses) * 0.40 + detention * 25  # Include detention time in the calculation
    return total_pay




import os
import sys
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def save_data(data):
    document = Document()

    # Set font and size
    font_name = "Times New Roman"
    font_size = 12

 # Add the truck image
    truck_image_path = 'C:/Users/14842/Documents/App/trucking.png'  
    document.add_picture(truck_image_path)

    # Create a title paragraph
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run("Driver's Week")
    title_run.bold = True
    title_run.font.name = font_name
    title_run.font.size = Pt(14)

    # Add driver information
    driver_paragraph = document.add_paragraph("Driver: ")
    driver_run = driver_paragraph.add_run(data['driver_name'])
    driver_run.bold = True
    driver_run.font.name = font_name
    driver_run.font.size = Pt(font_size)

    document.add_paragraph("_________________________________________________________________________________________________________")

    # Create a table for trips
    trips_table = document.add_table(rows=1, cols=3)
    trips_table.autofit = False
    trips_table.columns[0].width = Cm(3)
    trips_table.columns[1].width = Cm(6)
    trips_table.columns[2].width = Cm(3)

    # Add table headers
    headers = trips_table.rows[0].cells
    headers[0].text = "TRIPS:"
    headers[1].text = "TRIP DETAILS:"
    headers[2].text = "TRIP DATE:"
    for cell in headers:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = font_name
        cell.paragraphs[0].runs[0].font.size = Pt(font_size)

    # Add trip data to the table
    for i, trip in enumerate(data['trips'], start=1):
        if trip['trip']:
            trip_row = trips_table.add_row().cells
            trip_row[0].text = f"Trip {i}: ${trip['trip']}"
            trip_row[1].text = trip['details']
            trip_row[2].text = trip['date']

    # Style the table cells
    for row in trips_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Add expenses section
    expenses_paragraph = document.add_paragraph("_________________________________________________________________________________________________________")
    expenses_paragraph_run = expenses_paragraph.add_run("Expenses:")
    expenses_paragraph_run.font.name = font_name
    expenses_paragraph_run.font.size = Pt(font_size)
    expenses_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    expenses_paragraph_run.bold = True

    diesel_paragraph = document.add_paragraph(f"Diesel: ${data['diesel']}")
    diesel_paragraph.runs[0].font.name = font_name
    diesel_paragraph.runs[0].font.size = Pt(font_size)

    tolls_paragraph = document.add_paragraph(f"Tolls: ${data['tolls']}")
    tolls_paragraph.runs[0].font.name = font_name
    tolls_paragraph.runs[0].font.size = Pt(font_size)

    document.add_paragraph("_________________________________________________________________________________________________________")

     # Add detention time information
    detention_input = data.get('detention')
    if detention_input.strip():
        detention = float(detention_input)
    else:
        detention = 0.0
    detention_pay = detention * 25  # Calculate the detention pay
    detention_paragraph = document.add_paragraph(f"Detention Time: ${detention_pay:.2f}")
    detention_paragraph.runs[0].font.name = font_name
    detention_paragraph.runs[0].font.size = Pt(font_size)
    

    document.add_paragraph("_________________________________________________________________________________________________________")

    # Calculate the pay with the actual detention time
    weekly_pay = calculate_pay(data)
    data['week_pay'] = weekly_pay

    # Add weekly pay and date
    weekly_pay_paragraph = document.add_paragraph()
    weekly_pay_run = weekly_pay_paragraph.add_run(f"Driver's Weekly Pay: ${weekly_pay:.2f}")
    weekly_pay_run.bold = True
    weekly_pay_run.font.name = font_name
    weekly_pay_run.font.size = Pt(font_size)
    weekly_pay_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    pay_date_paragraph = document.add_paragraph(f"Pay Date: {data['final_date']}")
    pay_date_paragraph.runs[0].font.name = font_name
    pay_date_paragraph.runs[0].font.size = Pt(font_size)
    pay_date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Get the directory of the executable file
    if getattr(sys, 'frozen', False):
        script_dir = os.path.dirname(sys.executable)
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))

    # Construct the file path
    file_path = os.path.join(script_dir, 'driver_data.docx')

    # Save the document to the specified file path
    document.save(file_path)

    # Open the document
    os.startfile(file_path)

    # Wait for the application to open
    sg.popup_quick_message("Opening the document... Please wait.")

def main():
    icon_path = r"C:/Users/14842/Documents/App/dist/truck2.ico"
    sg.set_options(icon=icon_path)  # Set the program icon

    header_layout = [
        [sg.Frame('', [[sg.Text('VASQUEZ', font=('Helvetica', 30, 'bold'), text_color='#333333')],
                       [sg.Text('TRUCKING LLC', font=('Helvetica', 30, 'bold'), text_color='#333333')],
                       [sg.Text('48 Winding Brook Dr. Sinking Spring, PA 19608', font=('Helvetica', 13), text_color='#333333')],
                       [sg.Text('484-269-6029', font=('Helvetica', 13), text_color='#333333')]],
                  background_color='#444444', pad=(0, 0))
         ]
    ]

    truck_image = sg.Image('C:/Users/14842/Documents/App/truck.png', size=(200, 125))


    trip_rows = [
        [
            sg.Text(f'Trip {i+1}'),
            sg.Input(size=(10, 1), key=f'-TRIP{i+1}-', enable_events=True),
            sg.Text('Trip Details'),
            sg.Input(size=(20, 1), key=f'-DETAILS{i+1}-'),
            sg.CalendarButton('Select Date', target=f'-DATE{i+1}-', key=f'-CALENDAR{i+1}-', format='%Y-%m-%d'),
            sg.Input(size=(10, 1), key=f'-DATE{i+1}-', enable_events=True, readonly=True)
        ]
        for i in range(15)  # Changed to 15 trips
    ]

    main_layout = [
    [sg.Column(header_layout, background_color='#444444'), truck_image],
    [sg.Text('Name'), sg.Combo(DRIVERS, key='-DRIVER-', enable_events=True, readonly=True)],
    *trip_rows,  # Expanded trip_rows to include all 15 trips
    [sg.Text('Diesel'), sg.Input(key='-DIESEL-', enable_events=True)],
    [sg.Text('Tolls'), sg.Input(key='-TOLLS-', enable_events=True)],
    [sg.Text('Detention Time (Hours)'), sg.Input(key='-DETENTION-', enable_events=True)],  # Added Detention Time field
    [sg.CalendarButton('Select Pay Date', target='-FINALDATE-', key='-CALENDAR-', format='%Y-%m-%d'),
     sg.Input(size=(10, 1), key='-FINALDATE-', enable_events=True, readonly=True)],
    [sg.Button('Calculate', bind_return_key=True), sg.Button('Clear'), sg.Button('Exit')]
]

    window = sg.Window('Week Pay', main_layout)

    while True:
        event, values = window.read()
        if event == sg.WINDOW_CLOSED or event == 'Exit':
            break
        elif event == 'Calculate':
            trip_dates = [values[f'-DATE{i+1}-'] for i in range(15)]  # Changed to 15 trips

            driver_data = {
                'driver_name': values['-DRIVER-'],
                'trips': [
                    {
                        'trip': values[f'-TRIP{i+1}-'],
                        'details': values[f'-DETAILS{i+1}-'],
                        'date': trip_dates[i]
                    }
                    for i in range(15)  # Changed to 15 trips
                ],
                'diesel': values['-DIESEL-'],
                'tolls': values['-TOLLS-'],
                'detention': values['-DETENTION-'],  # Include the detention value in driver_data
                'final_date': values['-FINALDATE-']  # Add the final date to the driver_data dictionary
            }

            pay = calculate_pay(driver_data)
            if pay is not None:
                driver_data['week_pay'] = pay  # Add the week pay to the driver_data dictionary
                sg.popup(f"Driver's Weekly Pay: ${pay:.2f}")

                save_data(driver_data)

    
                
        elif event == 'Clear':
            # Clear the input fields
            for i in range(15):  # Changed to 15 trips
                window[f'-TRIP{i+1}-'].update('')
                window[f'-DETAILS{i+1}-'].update('')
                window[f'-DATE{i+1}-'].update('')
            window['-DRIVER-'].update('')
            window['-DIESEL-'].update('')
            window['-TOLLS-'].update('')
            window['-FINALDATE-'].update('')
            window['-DETENTION-'].update('')
            window[f'-TRIP1-'].set_focus()

    window.close()


if __name__ == "__main__":
    main()





